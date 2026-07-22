from pathlib import Path
import re
import textwrap

from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image as RLImage,
    PageBreak,
)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "generated"
CODE_SHOTS = ASSETS / "code"
DIAGRAMS = ASSETS / "diagrams"


PROJECT = {
    "nama": "Manhattan",
    "studi_kasus": "Aplikasi real-time group chat dengan end-to-end encryption.",
    "target_pengguna": "Pengguna yang membutuhkan ruang chat sementara, cepat, dan privat melalui browser.",
    "permasalahan": "Chat biasa sering menyimpan plaintext di server, sulit menjaga privasi, dan rawan bentrok sesi ketika beberapa client aktif.",
    "solusi": "Manhattan memakai WebSocket/STOMP untuk komunikasi real-time, Web Crypto API untuk RSA/AES di browser, dan Spring Boot hanya meneruskan ciphertext serta mengelola room/session.",
}

FUNCTIONAL_REQUIREMENTS = [
    ("FR-01", "User dapat membuat room baru dengan nama alphanumeric 3-15 karakter."),
    ("FR-02", "User dapat bergabung ke room yang sudah ada dengan password jika room diproteksi."),
    ("FR-03", "Sistem memvalidasi satu sesi aktif per IP address pada WebSocket handshake."),
    ("FR-04", "Client mengenkripsi pesan dengan AES-256-GCM sebelum dikirim ke server."),
    ("FR-05", "Client bertukar AES key secara aman menggunakan RSA public key antar peserta."),
    ("FR-06", "Server meneruskan ciphertext ke topic room tanpa membaca plaintext pesan."),
    ("FR-07", "Sistem membatasi percobaan password salah sebanyak 5 kali lalu lockout 60 detik."),
    ("FR-08", "User dapat melihat daftar peserta aktif di room."),
]

LAYER_TABLE = [
    ("Controller", "RoomController, MessageController, KeyExchangeController", "Menerima STOMP message dari browser dan mengirim response/event."),
    ("Service", "RoomService, SessionService, RateLimitService, MessageRelayService, KeyExchangeService", "Menjalankan validasi, session lifecycle, rate limit, relay pesan, dan key exchange."),
    ("Repository", "RoomRepository, SessionRepository, QueuedMessageRepository, RateLimitRepository", "Akses data via Spring Data JPA."),
    ("Entity", "Room, Session, QueuedMessage, RateLimit, SessionStatus", "Representasi tabel database dan status domain."),
    ("DTO", "RoomCreationResult, RoomJoinResult, RoomInfo, ParticipantInfo", "Payload hasil operasi yang dikirim antar layer."),
]

API_ROWS = [
    ("WebSocket CONNECT", "/ws", "Membuka koneksi STOMP native WebSocket."),
    ("WebSocket CONNECT", "/ws-sockjs", "Endpoint fallback SockJS untuk browser lama."),
    ("STOMP SEND", "/app/room.create", "Membuat room, membuat session creator, broadcast USER_JOINED."),
    ("STOMP SEND", "/app/room.join", "Join room, verifikasi password, reset/record rate limit, broadcast participant."),
    ("STOMP SEND", "/app/room.info", "Mengambil info room: password, jumlah peserta, status aktif."),
    ("STOMP SEND", "/app/message.send", "Relay ciphertext dan IV ke topic room."),
    ("STOMP SEND", "/app/key.exchange", "Forward encrypted AES key ke private queue target."),
    ("STOMP SUBSCRIBE", "/topic/room/{roomName}", "Menerima pesan terenkripsi room."),
    ("STOMP SUBSCRIBE", "/topic/room/{roomName}/events", "Menerima event room dan daftar peserta."),
    ("STOMP SUBSCRIBE", "/user/queue/private", "Menerima response private, error, dan AES key exchange."),
]

TEAM_ROWS = [
    (
        "Richo Arbianto",
        "Backend - Room & Session Management",
        "RoomController.java, RoomService.java, RoomRepository.java, Room.java, Session.java, SessionService.java",
        "Sinkronisasi state sesi saat client disconnect tiba-tiba.",
    ),
    (
        "Richo Arbianto",
        "Backend - Security & Messaging",
        "IpGuardInterceptor.java, RateLimitService.java, MessageController.java, MessageRelayService.java, KeyExchangeController.java",
        "Implementasi verifikasi Argon2id PHC format di server-side.",
    ),
    (
        "Mohammad Saif Al-Islam",
        "Frontend - UI & Crypto",
        "crypto.js, keystore.js, argon2.js, ui/chat-interface.js, ui/room-entry.js",
        "Web Crypto API tidak sinkron, semua operasi harus async/await.",
    ),
    (
        "Mohammad Saif Al-Islam",
        "Frontend - Networking & Testing",
        "websocket-client.js, key-exchange.js, room-controller.js, semua file *.test.js",
        "STOMP reconnect logic dan race condition saat key exchange.",
    ),
]

DOC_CHECKLIST = [
    ("README", "Ada", "README.md berisi overview, cara kerja, tech stack, setup lokal, testing, deploy, dan batasan desain."),
    ("API Documentation", "Dibuat", "docs/API.pdf"),
    ("User Manual", "Dibuat", "docs/UserManual.pdf"),
    ("Installation Guide", "Dibuat", "docs/InstallationGuide.pdf"),
    ("Testing Report", "Dibuat", "docs/TestingReport.pdf"),
    ("Functional Requirement", "Dibuat", "docs/FunctionalRequirement.pdf"),
    ("ERD", "Dibuat", "docs/ERD.pdf"),
    ("Class Diagram", "Dibuat", "docs/ClassDiagram.pdf"),
    ("Architecture Diagram", "Dibuat", "docs/ArchitectureDiagram.pdf"),
    ("Deployment Report", "Dibuat", "docs/DeploymentReport.pdf"),
]


def ensure_dirs():
    CODE_SHOTS.mkdir(parents=True, exist_ok=True)
    DIAGRAMS.mkdir(parents=True, exist_ok=True)


def read_lines(rel_path, start, end):
    path = ROOT / rel_path
    lines = path.read_text().splitlines()
    selected = lines[start - 1:end]
    return "\n".join(f"{i:>4}  {line}" for i, line in zip(range(start, end + 1), selected))


def font(size=18, bold=False):
    candidates = [
        "/System/Library/Fonts/Menlo.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def text_font(size=24, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def code_shot(name, title, rel_path, start, end):
    code = read_lines(rel_path, start, end)
    mono = font(17)
    title_font = text_font(24, True)
    line_h = 24
    lines = code.splitlines()
    width = max(980, max(int(mono.getlength(line)) for line in lines) + 64)
    height = 78 + len(lines) * line_h + 34
    img = Image.new("RGB", (width, height), "#0b1020")
    d = ImageDraw.Draw(img)
    d.rectangle([0, 0, width, 54], fill="#1f3a5f")
    d.text((24, 14), title, fill="#ffffff", font=title_font)
    y = 74
    for line in lines:
        line_no = line[:6]
        body = line[6:]
        d.text((24, y), line_no, fill="#7dd3fc", font=mono)
        d.text((92, y), body, fill="#e5e7eb", font=mono)
        y += line_h
    out = CODE_SHOTS / f"{name}.png"
    img.save(out)
    return out


def rounded_rect(draw, xy, fill, outline="#1f2937", width=2, radius=12):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_arrow(draw, start, end, color="#0f172a", width=3):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - sign * 14, y2 - 8), (x2 - sign * 14, y2 + 8)]
    else:
        sign = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 8, y2 - sign * 14), (x2 + 8, y2 - sign * 14)]
    draw.polygon(pts, fill=color)


def label(draw, box, title, subtitle=""):
    x1, y1, x2, y2 = box
    tf = text_font(23, True)
    sf = text_font(17)
    tw = draw.textlength(title, font=tf)
    draw.text((x1 + (x2 - x1 - tw) / 2, y1 + 20), title, fill="#0f172a", font=tf)
    if subtitle:
        wrapped = textwrap.wrap(subtitle, width=26)
        y = y1 + 56
        for line in wrapped[:3]:
            lw = draw.textlength(line, font=sf)
            draw.text((x1 + (x2 - x1 - lw) / 2, y), line, fill="#334155", font=sf)
            y += 22


def architecture_diagram():
    img = Image.new("RGB", (1500, 780), "#f8fafc")
    d = ImageDraw.Draw(img)
    title = text_font(38, True)
    d.text((60, 42), "Architecture Diagram - Manhattan", fill="#0b2545", font=title)
    boxes = [
        ("Browser", "Vanilla JS, Web Crypto, STOMP Client", (60, 180, 310, 315), "#dbeafe"),
        ("Controller", "Room/Message/KeyExchange Controller", (380, 180, 650, 315), "#e0f2fe"),
        ("Service", "Room, Session, RateLimit, Relay, KeyExchange", (720, 180, 1010, 315), "#dcfce7"),
        ("Repository", "Spring Data JPA Repository", (1080, 180, 1320, 315), "#fef9c3"),
        ("Database", "MySQL: rooms, sessions, queue, rate_limits", (1120, 455, 1390, 600), "#fee2e2"),
    ]
    for title_txt, sub, box, fill in boxes:
        rounded_rect(d, box, fill)
        label(d, box, title_txt, sub)
    draw_arrow(d, (310, 248), (380, 248))
    draw_arrow(d, (650, 248), (720, 248))
    draw_arrow(d, (1010, 248), (1080, 248))
    draw_arrow(d, (1200, 315), (1200, 455))
    draw_arrow(d, (390, 315), (255, 455))
    rounded_rect(d, (60, 455, 450, 620), "#ffffff")
    label(d, (60, 455, 450, 620), "Response/Event Flow", "topic room, room events, user private queue")
    draw_arrow(d, (255, 455), (255, 315))
    out = DIAGRAMS / "architecture.png"
    img.save(out)
    return out


def erd_diagram():
    img = Image.new("RGB", (1500, 1000), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((60, 42), "ERD - Manhattan Database", fill="#0b2545", font=text_font(38, True))
    boxes = {
        "rooms": (80, 150, 490, 430),
        "sessions": (670, 130, 1130, 460),
        "message_queue": (670, 570, 1130, 880),
        "rate_limits": (80, 570, 490, 840),
    }
    fields = {
        "rooms": ["PK id BIGINT", "UQ name VARCHAR(15)", "password_hash VARCHAR(255)", "creator_ip VARCHAR(45)", "created_at TIMESTAMP", "is_active BOOLEAN"],
        "sessions": ["PK id BIGINT", "ip_address VARCHAR(45)", "FK room_name -> rooms.name", "stomp_session_id VARCHAR(64)", "display_name VARCHAR(45)", "connected_at TIMESTAMP", "disconnected_at TIMESTAMP", "last_activity_at TIMESTAMP", "status ENUM"],
        "message_queue": ["PK id BIGINT", "target_ip VARCHAR(45)", "FK room_name -> rooms.name", "sender_ip VARCHAR(45)", "ciphertext BLOB", "iv VARBINARY(16)", "created_at TIMESTAMP"],
        "rate_limits": ["PK id BIGINT", "client_ip VARCHAR(45)", "room_name VARCHAR(15)", "failed_attempts INT", "locked_until TIMESTAMP", "last_attempt_at TIMESTAMP", "UQ client_ip + room_name"],
    }
    for name, box in boxes.items():
        rounded_rect(d, box, "#f8fafc", "#334155")
        x1, y1, x2, y2 = box
        d.rectangle([x1, y1, x2, y1 + 48], fill="#1f3a5f")
        d.text((x1 + 16, y1 + 12), name, fill="#ffffff", font=text_font(24, True))
        y = y1 + 64
        for f in fields[name]:
            d.text((x1 + 18, y), f, fill="#111827", font=text_font(18))
            y += 31
    draw_arrow(d, (490, 260), (670, 260))
    d.text((530, 226), "1:N", fill="#0f172a", font=text_font(22, True))
    draw_arrow(d, (790, 570), (790, 460))
    d.text((808, 500), "1:N", fill="#0f172a", font=text_font(22, True))
    d.text((120, 900), "Catatan: rate_limits memakai room_name untuk pasangan unik client_ip + room_name, tetapi schema tidak mendefinisikan FK eksplisit.", fill="#475569", font=text_font(20))
    out = DIAGRAMS / "erd.png"
    img.save(out)
    return out


def class_diagram():
    img = Image.new("RGB", (1600, 1050), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((60, 42), "Class Diagram - Backend Manhattan", fill="#0b2545", font=text_font(38, True))
    classes = {
        "RoomController": (60, 140, 410, 330, ["- roomService", "- sessionService", "- rateLimitService", "+ createRoom()", "+ joinRoom()", "+ getRoomInfo()"]),
        "RoomService": (500, 140, 850, 360, ["- roomRepository", "- sessionRepository", "+ createRoom()", "+ joinRoom()", "+ getRoomInfo()", "+ verifyArgon2Password()"]),
        "RoomRepository": (960, 150, 1300, 300, ["extends JpaRepository", "+ findByName()", "+ existsByName()"]),
        "Room": (960, 430, 1300, 650, ["- id: Long", "- name: String", "- passwordHash: String", "- creatorIp: String", "- isActive: boolean"]),
        "SessionService": (500, 470, 850, 690, ["- sessionRepository", "+ createSession()", "+ markDisconnected()", "+ hasActiveSession()", "+ releaseIp()"]),
        "SessionRepository": (60, 480, 410, 660, ["extends JpaRepository", "+ findByIpAddressAndStatus()", "+ findByRoomNameAndStatus()", "+ countByRoomNameAndStatus()"]),
        "Session": (60, 760, 410, 985, ["- id: Long", "- ipAddress: String", "- roomName: String", "- stompSessionId: String", "- status: SessionStatus"]),
        "DTO Result": (500, 780, 850, 960, ["RoomCreationResult", "RoomJoinResult", "RoomInfo", "ParticipantInfo"]),
    }
    for name, (x1, y1, x2, y2, items) in classes.items():
        rounded_rect(d, (x1, y1, x2, y2), "#f8fafc", "#334155")
        d.rectangle([x1, y1, x2, y1 + 44], fill="#1f3a5f")
        d.text((x1 + 14, y1 + 10), name, fill="#ffffff", font=text_font(22, True))
        y = y1 + 60
        for item in items:
            d.text((x1 + 16, y), item, fill="#111827", font=text_font(17))
            y += 27
    draw_arrow(d, (410, 235), (500, 235))
    draw_arrow(d, (850, 235), (960, 235))
    draw_arrow(d, (1130, 300), (1130, 430))
    draw_arrow(d, (675, 360), (675, 470))
    draw_arrow(d, (500, 580), (410, 580))
    draw_arrow(d, (235, 660), (235, 760))
    d.text((430, 214), "association", fill="#475569", font=text_font(16))
    d.text((875, 214), "association", fill="#475569", font=text_font(16))
    d.text((1030, 360), "maps entity", fill="#475569", font=text_font(16))
    d.text((420, 900), "DTO dipakai sebagai object hasil operasi service.", fill="#475569", font=text_font(20))
    out = DIAGRAMS / "class_diagram.png"
    img.save(out)
    return out


def add_page_number(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 9)
    canvas.setFillColor(colors.HexColor("#64748B"))
    canvas.drawRightString(A4[0] - 0.55 * inch, 0.42 * inch, f"Halaman {doc.page}")
    canvas.restoreState()


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="DocTitle", fontName="Helvetica-Bold", fontSize=22, leading=27, textColor=colors.HexColor("#0B2545"), spaceAfter=12))
    styles.add(ParagraphStyle(name="H1x", fontName="Helvetica-Bold", fontSize=15, leading=19, textColor=colors.HexColor("#1F3A5F"), spaceBefore=14, spaceAfter=8))
    styles.add(ParagraphStyle(name="Bodyx", fontName="Helvetica", fontSize=10.5, leading=14, spaceAfter=7))
    styles.add(ParagraphStyle(name="Smallx", fontName="Helvetica", fontSize=8.8, leading=11, textColor=colors.HexColor("#475569"), spaceAfter=6))
    return styles


def table(data, widths=None, font_size=8.8):
    t = Table(data, colWidths=widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F3A5F")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), font_size),
        ("LEADING", (0, 0), (-1, -1), font_size + 2),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EFF6FF")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return t


def img_flowable(path, width=6.7 * inch, max_height=7.15 * inch):
    im = Image.open(path)
    w, h = im.size
    scale = min(width / w, max_height / h)
    return RLImage(str(path), width=w * scale, height=h * scale)


def build_pdf(filename, title, sections):
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(DOCS / filename),
        pagesize=A4,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.6 * inch,
    )
    story = [Paragraph(title, styles["DocTitle"])]
    for kind, payload in sections:
        if kind == "h1":
            story.append(Paragraph(payload, styles["H1x"]))
        elif kind == "p":
            story.append(Paragraph(payload, styles["Bodyx"]))
        elif kind == "small":
            story.append(Paragraph(payload, styles["Smallx"]))
        elif kind == "table":
            story.append(table(*payload))
            story.append(Spacer(1, 10))
        elif kind == "image":
            story.append(img_flowable(payload))
            story.append(Spacer(1, 10))
        elif kind == "break":
            story.append(PageBreak())
    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)


def set_cell(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.bold = bold
    if bold:
        r.font.color.rgb = RGBColor(255, 255, 255)
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_docx_table(doc, rows):
    t = doc.add_table(rows=1, cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, head in enumerate(rows[0]):
        set_cell(t.rows[0].cells[i], head, True, "1F3A5F")
    for row in rows[1:]:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], str(val))
    doc.add_paragraph()
    return t


def add_docx_image(doc, path, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
    return p


def create_docx(shots, diagrams):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Dokumentasi Software Engineering - Manhattan")
    r.font.name = "Arial"
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0, 0, 0)
    add_para(doc, "Nama: Richo Arbianto | NIM: 41524010083 | Mata kuliah: Pemrograman Berorientasi Objek")

    add_heading(doc, "1. Software Requirement")
    add_docx_table(doc, [["Item", "Isi"]] + [[k.replace("_", " ").title(), v] for k, v in PROJECT.items()])
    add_docx_table(doc, [["ID", "Functional Requirement"]] + FUNCTIONAL_REQUIREMENTS)
    add_para(doc, "Screenshot kode pendukung requirement:")
    add_docx_image(doc, shots["req_room_controller"], 6.2)

    add_heading(doc, "2. Software Architecture")
    add_docx_image(doc, diagrams["architecture"], 6.3)
    add_docx_table(doc, [["Layer", "Class/File", "Tanggung Jawab"]] + LAYER_TABLE)
    add_para(doc, "Arsitektur layered dipilih karena memisahkan input WebSocket, aturan bisnis, akses data, dan model database. Pemisahan ini membuat testing lebih mudah, perubahan repository tidak langsung merusak controller, dan logic keamanan seperti rate limit/session guard tetap terpusat di service.")
    add_docx_image(doc, shots["architecture_config"], 6.2)

    add_heading(doc, "3. Database Design")
    add_docx_image(doc, diagrams["erd"], 6.4)
    add_para(doc, "Entity utama: rooms menyimpan room, sessions menyimpan koneksi user aktif/nonaktif, message_queue menyimpan ciphertext sementara untuk user offline, rate_limits menyimpan percobaan password per IP dan room.")
    add_para(doc, "Primary key: id pada setiap tabel. Foreign key: sessions.room_name dan message_queue.room_name mengarah ke rooms.name dengan ON DELETE CASCADE. Relationship: satu room memiliki banyak session dan banyak queued message. rate_limits berelasi logis ke room_name melalui unique client_ip + room_name.")
    add_docx_image(doc, shots["database_schema"], 6.2)

    add_heading(doc, "4. Class Diagram")
    add_docx_image(doc, diagrams["class"], 6.4)
    add_para(doc, "Association terjadi antara controller-service, service-repository, dan repository-entity. Inheritance digunakan oleh repository interface yang extends JpaRepository. Composition/agregasi ringan terlihat pada service yang memiliki dependency repository melalui constructor injection.")
    add_docx_image(doc, shots["class_service"], 6.2)

    add_heading(doc, "5. REST API / WebSocket API")
    add_para(doc, "Project ini memakai WebSocket/STOMP, bukan REST endpoint klasik. Tabel berikut memakai Method untuk jenis operasi protokol.")
    add_docx_table(doc, [["Method", "Endpoint", "Fungsi"]] + API_ROWS)
    add_para(doc, "Endpoint pilihan: /app/room.join. Method: STOMP SEND. Request: { roomName, password, rsaPublicKey, displayName }. Response sukses dikirim ke /user/queue/private dengan type ROOM_JOINED, roomName, clientIp, participantCount. DTO/service: RoomJoinResult dari RoomService.joinRoom(), lalu SessionService.createSession(), RateLimitService.resetAttempts(), dan KeyExchangeService.broadcastPublicKey().")
    add_docx_image(doc, shots["api_room_join"], 6.2)

    add_heading(doc, "6. Dokumentasi Software Engineering")
    add_docx_table(doc, [["Dokumen", "Status", "Keterangan"]] + DOC_CHECKLIST)
    add_docx_image(doc, shots["doc_readme"], 6.2)

    doc.add_page_break()
    add_heading(doc, "7. Refleksi Individu")
    add_para(doc, "Kontribusi terbesar Richo Arbianto adalah membangun backend room/session/security/messaging: validasi room, lifecycle session, satu IP satu sesi, rate limiting, relay ciphertext, key exchange routing, entity/repository, dan schema MySQL.")
    add_docx_table(doc, [["Nama Anggota", "Peran / Modul", "File / Component", "Kendala Teknis"]] + TEAM_ROWS)
    add_para(doc, "Error paling sulit untuk Richo adalah sinkronisasi state sesi saat client disconnect tiba-tiba. Masalahnya: IP bisa tetap dianggap aktif jika WebSocket close tidak membersihkan session, sehingga user yang sama gagal reconnect. Solusinya dibuat melalui IpGuardInterceptor untuk menyimpan clientIp saat handshake dan SessionService untuk markDisconnected/releaseIp saat disconnect.")
    add_para(doc, "Kendala kedua adalah verifikasi Argon2id PHC format di server-side. Hash dibuat di client, tetapi server harus parse parameter m/t/p, decode salt/hash base64 tanpa padding, lalu membandingkan hasil re-hash secara constant-time.")
    add_para(doc, "Fitur satu tahun lagi: persistent encrypted message history, file sharing terenkripsi, invite link dengan expiry, moderation room, multi-device identity, push notification, dan dashboard observability untuk koneksi WebSocket.")
    add_docx_image(doc, shots["reflection_ipguard"], 6.2)
    add_docx_image(doc, shots["reflection_argon2"], 6.2)

    out = DOCS / "SoftwareEngineeringReport.docx"
    doc.save(out)
    return out


def main():
    ensure_dirs()
    shots = {
        "req_room_controller": code_shot("01_requirement_room_controller", "Requirement Evidence - Client RoomController", "client/src/room-controller.js", 39, 93),
        "architecture_config": code_shot("02_architecture_websocket_config", "Architecture Evidence - WebSocketConfig", "server/src/main/java/com/manhattan/config/WebSocketConfig.java", 17, 51),
        "database_schema": code_shot("03_database_schema", "Database Evidence - schema.sql", "server/src/main/resources/schema.sql", 1, 40),
        "class_service": code_shot("04_class_room_service", "Class Evidence - RoomService Dependencies", "server/src/main/java/com/manhattan/service/RoomService.java", 18, 86),
        "api_room_join": code_shot("05_api_room_join", "API Evidence - /app/room.join", "server/src/main/java/com/manhattan/controller/RoomController.java", 88, 140),
        "doc_readme": code_shot("06_documentation_readme", "Documentation Evidence - README", "README.md", 73, 139),
        "reflection_ipguard": code_shot("07_reflection_ipguard", "Reflection Evidence - IpGuardInterceptor", "server/src/main/java/com/manhattan/interceptor/IpGuardInterceptor.java", 31, 58),
        "reflection_argon2": code_shot("08_reflection_argon2", "Reflection Evidence - Argon2 PHC Verify", "server/src/main/java/com/manhattan/service/RoomService.java", 151, 191),
        "controller_message": code_shot("09_controller_message", "Controller Evidence - MessageController", "server/src/main/java/com/manhattan/controller/MessageController.java", 25, 65),
        "repository": code_shot("10_repository_session", "Repository Evidence - SessionRepository", "server/src/main/java/com/manhattan/repository/SessionRepository.java", 1, 18),
        "dto": code_shot("11_dto_room_join_result", "DTO Evidence - RoomJoinResult", "server/src/main/java/com/manhattan/dto/RoomJoinResult.java", 1, 27),
    }
    diagrams = {
        "architecture": architecture_diagram(),
        "erd": erd_diagram(),
        "class": class_diagram(),
    }

    build_pdf("FunctionalRequirement.pdf", "Functional Requirement - Manhattan", [
        ("h1", "Software Requirement"),
        ("table", ([["Item", "Isi"]] + [[k.replace("_", " ").title(), v] for k, v in PROJECT.items()], [1.65 * inch, 5.0 * inch], 9)),
        ("h1", "Functional Requirement"),
        ("table", ([["ID", "Functional Requirement"]] + FUNCTIONAL_REQUIREMENTS, [0.8 * inch, 5.85 * inch], 9)),
        ("h1", "Screenshot Kode"),
        ("image", shots["req_room_controller"]),
    ])

    build_pdf("ArchitectureDiagram.pdf", "Software Architecture - Manhattan", [
        ("image", diagrams["architecture"]),
        ("h1", "Tabel Layer"),
        ("table", ([["Layer", "Class/File", "Tanggung Jawab"]] + LAYER_TABLE, [1.0 * inch, 2.5 * inch, 3.15 * inch], 8.2)),
        ("p", "Arsitektur layered dipilih karena memisahkan controller, service, repository, entity, dan DTO. Dengan pola ini, logic bisnis tidak bercampur dengan detail transport WebSocket atau akses database."),
        ("h1", "Screenshot Kode"),
        ("image", shots["architecture_config"]),
    ])

    build_pdf("ERD.pdf", "Database Design - Manhattan", [
        ("image", diagrams["erd"]),
        ("h1", "Entity, Key, dan Relationship"),
        ("p", "Entity utama: rooms, sessions, message_queue, dan rate_limits. Primary key setiap tabel adalah id. Foreign key eksplisit: sessions.room_name -> rooms.name dan message_queue.room_name -> rooms.name. Relationship utama: rooms 1:N sessions dan rooms 1:N message_queue."),
        ("h1", "Screenshot Kode"),
        ("image", shots["database_schema"]),
    ])

    build_pdf("ClassDiagram.pdf", "Class Diagram - Manhattan", [
        ("image", diagrams["class"]),
        ("h1", "Hubungan Class"),
        ("p", "Association: controller memakai service, service memakai repository, repository mengakses entity. Inheritance: repository extends JpaRepository. Composition/agregasi: service menyimpan dependency repository melalui constructor injection."),
        ("h1", "Screenshot Kode"),
        ("image", shots["class_service"]),
        ("image", shots["repository"]),
        ("image", shots["dto"]),
    ])

    build_pdf("API.pdf", "API Documentation - Manhattan", [
        ("h1", "Endpoint WebSocket/STOMP"),
        ("table", ([["Method", "Endpoint", "Fungsi"]] + API_ROWS, [1.35 * inch, 2.05 * inch, 3.25 * inch], 8)),
        ("h1", "Endpoint Detail: /app/room.join"),
        ("p", "Method: STOMP SEND. Request: roomName, password, rsaPublicKey, displayName. Response sukses: ROOM_JOINED ke /user/queue/private. DTO: RoomJoinResult. Service: RoomService.joinRoom(), RateLimitService, SessionService, KeyExchangeService."),
        ("h1", "Screenshot Kode"),
        ("image", shots["api_room_join"]),
        ("image", shots["controller_message"]),
    ])

    build_pdf("UserManual.pdf", "User Manual - Manhattan", [
        ("h1", "Panduan Pengguna"),
        ("p", "1. Buka aplikasi di browser. 2. Masukkan display name dan nama room. 3. Pilih create room atau join room. 4. Jika room memiliki password, masukkan password. 5. Setelah masuk room, kirim pesan melalui input chat. 6. Pesan akan terenkripsi dari browser sebelum dikirim."),
        ("image", DOCS / "screenshots" / "room-entry.png"),
        ("image", DOCS / "screenshots" / "chat.png"),
        ("h1", "Screenshot Kode"),
        ("image", shots["req_room_controller"]),
    ])

    build_pdf("InstallationGuide.pdf", "Installation Guide - Manhattan", [
        ("h1", "Prerequisite"),
        ("p", "Java 23, Node.js 20 LTS, MySQL 8.0, Gradle wrapper pada folder server."),
        ("h1", "Langkah Instalasi"),
        ("p", "1. Buat database manhattan. 2. Jalankan schema.sql. 3. Jalankan server dengan cd server && ./gradlew bootRun. 4. Jalankan client dengan cd client && npm install && npm run dev. 5. Buka http://localhost:3000."),
        ("h1", "Screenshot Kode"),
        ("image", shots["doc_readme"]),
    ])

    build_pdf("TestingReport.pdf", "Testing Report - Manhattan", [
        ("h1", "Strategi Testing"),
        ("p", "Server diuji dengan JUnit 5, jqwik property-based tests, H2, dan Testcontainers. Client diuji dengan Jest, fast-check, dan Playwright untuk e2e."),
        ("h1", "Catatan Hasil"),
        ("p", "Di repo terdapat playwright-report yang menunjukkan kegagalan e2e karena executable Chromium Playwright belum terpasang di cache lokal. Perbaikannya adalah menjalankan npx playwright install sebelum npm run test:e2e."),
        ("h1", "Screenshot Kode"),
        ("image", shots["doc_readme"]),
    ])

    build_pdf("DeploymentReport.pdf", "Deployment Report - Manhattan", [
        ("h1", "Deployment"),
        ("p", "Deployment production memakai Ubuntu 22.04, Nginx reverse proxy, MySQL, systemd service, build Gradle untuk backend, dan build Tailwind/esbuild untuk frontend. Script deploy.sh mengotomasi instalasi dependency, build, konfigurasi Nginx SSL, dan service startup."),
        ("h1", "Dokumen Pendukung"),
        ("p", "File pendukung deployment: deploy.sh, deploy/manhattan.conf, deploy/manhattan.service, application-prod.yml."),
        ("h1", "Screenshot Kode"),
        ("image", shots["doc_readme"]),
    ])

    create_docx(shots, diagrams)
    print("Generated documentation in docs/")


if __name__ == "__main__":
    main()
